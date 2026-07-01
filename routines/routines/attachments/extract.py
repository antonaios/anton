"""Local-only text extraction for chat document-attachments (#chat-attachments).

SECURITY INVARIANT (§5.2 / CLAUDE.md #no-mnpi-to-cloud): every extractor here is
LOCAL — pure-CPU for the Office/text formats, and the LOCAL Ollama lane (never a
cloud model) for the scanned-PDF fallback. The uploaded bytes are read off local
disk and handed ONLY to these local parsers; they are NEVER placed in a prompt,
a chat history, or any cloud request. Only the extracted TEXT this module
returns is ever injected into a chat turn.

Per format (all local):
  * PDF   — text-first via ``routines.intake.pdf_render.extract_text_per_page``
            (pypdfium2, NO LLM). When the average per-page text is too sparse
            (a scanned / image-only PDF) we fall back to
            ``routines.intake.parse.ingest_pdf(..., client=OllamaClient())``,
            which pins the multimodal lane to LOCAL Ollama (gemma) — the image
            path. Never cloud.
  * .pptx — ``routines.shared.pptx_to_markdown.pptx_to_markdown(
            describe_images=False)`` — pure-CPU python-pptx text/table/chart
            extraction, no LLM touched.
  * .docx — python-docx, pure-CPU (paragraphs + tables).
  * .xlsx — openpyxl read-only, pure-CPU (per-sheet cell grid → text).
  * .txt / .md / .csv — direct UTF-8 read (latin-1 fallback).

The extracted text is capped at ``MAX_EXTRACTED_CHARS`` (truncate + flag) so a
huge document can't blow the chat context window.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


# Cap the extracted text so a giant document can't blow the chat context window.
# ~24k chars ≈ 6k tokens — generous headroom for a teaser / a few pages of a CIM
# while staying well under qwen3:14b's window.
MAX_EXTRACTED_CHARS = 24_000

# Average non-empty per-page text length below which a PDF is treated as
# scanned / image-only and routed to the LOCAL image path. Mirrors the intake
# orchestrator's ``DEFAULT_TEXT_THRESHOLD`` rationale (150 chars/page).
_PDF_TEXT_THRESHOLD = 150

# Supported lowercase extensions (incl. the leading dot). The route rejects
# anything not in this set with a 422 BEFORE saving / extracting.
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
    {".pdf", ".docx", ".pptx", ".xlsx", ".txt", ".md", ".csv"}
)

# Plain-text extensions read directly as UTF-8.
_PLAINTEXT_EXTENSIONS: frozenset[str] = frozenset({".txt", ".md", ".csv"})


class AttachmentExtractError(RuntimeError):
    """Raised when a supported attachment cannot be extracted (corrupt file,
    parser failure). The route maps this to a 422."""


class UnsupportedAttachment(AttachmentExtractError):
    """Raised when the extension is not in ``SUPPORTED_EXTENSIONS``."""


@dataclass(frozen=True)
class ExtractedDoc:
    """Result of a LOCAL extraction.

    ``lane`` records which local path produced the text — for the audit /
    observability only ("local-cpu" for the Office/text parsers, "local-ollama"
    for the scanned-PDF image fallback). It is NEVER "cloud" — that would be a
    §5.2 violation.
    """

    text: str
    chars: int          # len(text) AFTER truncation
    truncated: bool
    lane: str           # "local-cpu" | "local-ollama"


def _truncate(text: str, lane: str) -> ExtractedDoc:
    """Cap ``text`` at ``MAX_EXTRACTED_CHARS`` (truncate + flag)."""
    truncated = len(text) > MAX_EXTRACTED_CHARS
    capped = text[:MAX_EXTRACTED_CHARS] if truncated else text
    return ExtractedDoc(
        text=capped, chars=len(capped), truncated=truncated, lane=lane,
    )


# ── public entry point ────────────────────────────────────────────────────────


def extract_text(path: Path) -> ExtractedDoc:
    """Extract text from ``path`` using ONLY local extractors, dispatched by
    extension. Raises ``UnsupportedAttachment`` for an unknown extension and
    ``AttachmentExtractError`` on a parser failure.

    The dispatch is purely local — NO cloud model is ever reached (§5.2)."""
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedAttachment(f"unsupported attachment type {ext!r}")

    if ext in _PLAINTEXT_EXTENSIONS:
        return _extract_plaintext(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".pptx":
        return _extract_pptx(path)
    if ext == ".xlsx":
        return _extract_xlsx(path)
    # Unreachable — SUPPORTED_EXTENSIONS and the branches above are in sync.
    raise UnsupportedAttachment(f"unsupported attachment type {ext!r}")


# ── plaintext (.txt / .md / .csv) ─────────────────────────────────────────────


def _extract_plaintext(path: Path) -> ExtractedDoc:
    """Direct read. UTF-8 first; latin-1 fallback so a non-UTF-8 file still
    yields text rather than erroring (the bytes are local-only either way)."""
    try:
        raw = path.read_bytes()
    except OSError as e:
        raise AttachmentExtractError(f"could not read {path.name}: {e}") from e
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")
    return _truncate(text, lane="local-cpu")


# ── PDF (local text-first, local-Ollama image fallback) ───────────────────────


def _extract_pdf(path: Path) -> ExtractedDoc:
    """Text-first via pypdfium2 (NO LLM); fall back to the LOCAL Ollama image
    path only for a scanned / image-only PDF. Never cloud."""
    from routines.intake.pdf_render import (
        PDFRenderError,
        extract_text_per_page,
    )

    try:
        pages = extract_text_per_page(path)
    except PDFRenderError as e:
        raise AttachmentExtractError(f"PDF read failed for {path.name}: {e}") from e

    if pages:
        avg_chars = sum(len(p.text) for p in pages) / max(len(pages), 1)
    else:
        avg_chars = 0.0

    if pages and avg_chars >= _PDF_TEXT_THRESHOLD:
        # Text-heavy PDF — assemble the per-page text directly. NO LLM.
        body = "\n\n".join(
            f"=== PAGE {p.page_number} ===\n{p.text}" for p in pages if p.text
        )
        return _truncate(body, lane="local-cpu")

    # Scanned / image-only (sparse extractable text) → the LOCAL image path.
    # ``ingest_pdf`` with an ``OllamaClient`` pins multimodal-extraction to the
    # LOCAL gemma lane (see routines/intake/parse.py) — the bytes never leave
    # the box. We force ``mode="image"`` so it doesn't re-extract text we
    # already know is sparse.
    return _extract_pdf_via_ollama(path)


def _extract_pdf_via_ollama(path: Path) -> ExtractedDoc:
    """Scanned-PDF fallback: render pages locally and read them with the LOCAL
    Ollama multimodal lane (``ingest_pdf`` → gemma). Returns the parsed
    document's operator-facing precis as the injected text — the local model's
    structured read of the image-only pages.

    SECURITY: ``OllamaClient()`` is a LOOPBACK client and ``ingest_pdf`` routes
    multimodal-extraction to the LOCAL lane for every sensitivity; the page
    images never reach a cloud model (§5.2)."""
    from routines.intake.parse import ingest_pdf
    from routines.intake.pdf_render import PDFRenderError
    from routines.shared.ollama_client import OllamaClient, OllamaError

    try:
        result = ingest_pdf(path, client=OllamaClient(), mode="image")
    except (PDFRenderError, OllamaError) as e:
        raise AttachmentExtractError(
            f"scanned-PDF local extraction failed for {path.name}: {e}"
        ) from e

    text = _parsed_document_to_text(result.parsed)
    return _truncate(text, lane="local-ollama")


def _parsed_document_to_text(parsed) -> str:  # noqa: ANN001 — intake.schema.ParsedDocument
    """Flatten the structured ``ParsedDocument`` from the image path into a
    plain-text block for chat injection. Best-effort and defensive — a missing
    field is simply skipped."""
    lines: list[str] = []
    summary = getattr(parsed, "summary", "") or ""
    if summary:
        lines.append(summary)
    highlights = getattr(parsed, "investment_highlights", None) or []
    if highlights:
        lines.append("")
        lines.append("Highlights:")
        lines.extend(f"- {h}" for h in highlights if h)
    financials = getattr(parsed, "financials", None) or []
    if financials:
        lines.append("")
        lines.append("Financials:")
        for f in financials:
            metric = getattr(f, "metric", "") or ""
            value = getattr(f, "value", "") or ""
            period = getattr(f, "period", "") or ""
            if metric or value:
                lines.append(f"- {metric}: {value} ({period})".strip())
    return "\n".join(lines).strip()


# ── .docx (python-docx, pure-CPU) ─────────────────────────────────────────────


def _extract_docx(path: Path) -> ExtractedDoc:
    """Extract paragraphs + table cells from a .docx via python-docx (no LLM)."""
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover — declared dep
        raise AttachmentExtractError("python-docx not installed") from e

    try:
        document = docx.Document(str(path))
    except Exception as e:  # noqa: BLE001 — any open/parse failure → clean 422
        raise AttachmentExtractError(f"could not open .docx {path.name}: {e}") from e

    blocks: list[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [(_cell_text(c)) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                blocks.append(line)

    return _truncate("\n".join(blocks), lane="local-cpu")


def _cell_text(cell) -> str:  # noqa: ANN001 — docx table cell
    try:
        return " ".join((cell.text or "").split())
    except Exception:  # noqa: BLE001 — one bad cell degrades to empty
        return ""


# ── .pptx (pptx_to_markdown, pure-CPU) ────────────────────────────────────────


def _extract_pptx(path: Path) -> ExtractedDoc:
    """Extract slide text/tables/charts via the pure-CPU pptx→markdown port.

    ``describe_images=False`` (CPU only — NO LLM touched, no image bytes leave
    the box)."""
    from routines.shared.pptx_to_markdown import (
        PptxConversionError,
        pptx_to_markdown,
    )

    try:
        md = pptx_to_markdown(path, describe_images=False)
    except PptxConversionError as e:
        raise AttachmentExtractError(
            f"could not convert .pptx {path.name}: {e}"
        ) from e
    return _truncate(md, lane="local-cpu")


# ── .xlsx (openpyxl read-only, pure-CPU) ──────────────────────────────────────


def _extract_xlsx(path: Path) -> ExtractedDoc:
    """Extract cell values from every sheet via openpyxl (read-only, no LLM)."""
    try:
        from openpyxl import load_workbook
    except ImportError as e:  # pragma: no cover — declared dep
        raise AttachmentExtractError("openpyxl not installed") from e

    try:
        # read_only + data_only: stream cells, return last-computed values
        # rather than formula strings; never executes anything.
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001 — any open/parse failure → clean 422
        raise AttachmentExtractError(f"could not open .xlsx {path.name}: {e}") from e

    blocks: list[str] = []
    try:
        for ws in wb.worksheets:
            blocks.append(f"## Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [_fmt_cell(v) for v in row]
                # Drop fully-empty rows.
                if any(c for c in cells):
                    blocks.append(" | ".join(cells).rstrip(" |"))
    finally:
        wb.close()

    return _truncate("\n".join(blocks), lane="local-cpu")


def _fmt_cell(value) -> str:  # noqa: ANN001 — openpyxl cell value
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)
