"""Deterministic text extraction for the digest crew (#ingest-digest stage 2).

PDF via pypdfium2's textpage API (the same path
``routines/intake/pdf_render.py::extract_text_per_page`` uses), DOCX via
python-docx (the same read ``routines/hinotes/processor.py`` uses), Markdown via
a plain UTF-8 read. Heavy deps are imported INSIDE the functions (deferred) so
this module stays importable — and file-path-loadable for the bridge test
suite — without them.

This is the reproducible parse layer (UA's deterministic half): no LLM, no
network. Output is bounded (file-size cap + page cap + char cap) so a hostile or
pathological doc can't exhaust memory the way the F-16 bitmap-bomb work guards
the render path.

INTEGRATION SEAM (#ingest-digest ↔ #32, surfaced 2026-06-13): the /triage crew
established that the shared crews venv has NO PDF libs and that a crew needing
PDFs should extract BRIDGE-SIDE (the routines venv already has pypdfium2 /
python-docx) and pass page-tagged text over the stdio boundary, rather than
adding deps to the operator-gated shared crews venv (see the crew-venv-no-pdf
project memory). This module is therefore destined to MOVE bridge-side at
integration; until then, if it runs in a crews venv without these deps it raises
a clear DigestExtractError telling the caller to extract bridge-side (MD still
works with no deps). The logic is unchanged by where it runs.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Bounds — generous for real deal docs (CIMs run a few MB / tens of pages),
# tight enough to neuter a pathological input. Mirror the intake caps' intent.
MAX_INPUT_BYTES = 100 * 1024 * 1024     # 100 MB
MAX_PAGES = 60                          # CIMs/IMs run long; cap the parse cost
MAX_CHARS = 400_000                     # ~100k tokens — the analyzer batches below this


class DigestExtractError(Exception):
    """Raised when a doc cannot be read at all (missing, oversized, corrupt).
    A legitimately EMPTY doc is NOT an error — it returns ``""`` (which the
    classifier treats as a fail-closed 'cannot verify public' signal).

    Messages are PATH-FREE and CONTENT-FREE (codex data-handling MED): a file
    name / path can encode a deal name, and these errors flow into
    ``DocAnalysis.error``. Callers identify the doc by its content hash, not by
    anything carried here."""


def _check_size(path: Path) -> None:
    try:
        size = path.stat().st_size
    except OSError:
        raise DigestExtractError("could not stat file") from None
    if size > MAX_INPUT_BYTES:
        raise DigestExtractError(
            f"file too large: {size} bytes exceeds the {MAX_INPUT_BYTES}-byte cap"
        )


def _bound(text: str, max_chars: int) -> str:
    return text if len(text) <= max_chars else text[:max_chars]


def _extract_pdf(path: Path, *, max_pages: int, max_chars: int) -> str:
    try:
        import pypdfium2 as pdfium
    except ImportError as e:  # the crews venv has no PDF lib by design (#32)
        raise DigestExtractError(
            "PDF extraction unavailable here — extract bridge-side (#32 pattern)"
        ) from e
    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception as e:  # noqa: BLE001
        raise DigestExtractError("PDF open failed") from e
    parts: list[str] = []
    try:
        try:
            for i in range(min(len(pdf), max_pages)):
                page = pdf[i]
                try:
                    tp = page.get_textpage()
                    try:
                        parts.append(tp.get_text_range() or "")
                    finally:
                        tp.close()
                finally:
                    page.close()
                # Stop early once we've gathered enough — no point parsing 60
                # pages of a doc we'll truncate at max_chars anyway.
                if sum(len(p) for p in parts) >= max_chars:
                    break
        except Exception as e:  # noqa: BLE001 — a bad page must fail THIS doc
            # (caught per-doc in analyze_one), not escape and abort the run
            # (codex-5.5 SEV-2). Path-free message.
            raise DigestExtractError("PDF extract failed") from e
    finally:
        pdf.close()
    return _bound("\n".join(parts), max_chars)


def _extract_docx(path: Path, *, max_chars: int) -> str:
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as e:  # the crews venv has no DOCX lib by design (#32)
        raise DigestExtractError(
            "DOCX extraction unavailable here — extract bridge-side (#32 pattern)"
        ) from e
    try:
        doc = Document(str(path))
    except Exception as e:  # noqa: BLE001
        raise DigestExtractError("DOCX open failed") from e
    try:
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as e:  # noqa: BLE001 — bad paragraph stream → one-doc error
        raise DigestExtractError("DOCX extract failed") from e
    return _bound(text, max_chars)


def _extract_md(path: Path, *, max_chars: int) -> str:
    try:
        return _bound(path.read_text(encoding="utf-8", errors="replace"), max_chars)
    except OSError:
        raise DigestExtractError("read failed") from None


def extract_text(
    path: str | Path,
    *,
    doc_type: str | None = None,
    max_pages: int = MAX_PAGES,
    max_chars: int = MAX_CHARS,
) -> str:
    """Extract bounded plain text from a PDF / DOCX / MD file.

    ``doc_type`` (``"pdf"``/``"docx"``/``"md"``) selects the path; when omitted
    it is inferred from the suffix. Returns ``""`` for an unsupported type or a
    legitimately empty doc. Raises :class:`DigestExtractError` only when the
    file cannot be read at all (missing / oversized / corrupt)."""
    path = Path(path)
    if not path.exists():
        raise DigestExtractError("not found")
    _check_size(path)

    kind = (doc_type or path.suffix.lower().lstrip(".")).lower()
    if kind in ("pdf",):
        return _extract_pdf(path, max_pages=max_pages, max_chars=max_chars)
    if kind in ("docx",):
        return _extract_docx(path, max_chars=max_chars)
    if kind in ("md", "markdown", "txt"):
        return _extract_md(path, max_chars=max_chars)
    logger.debug("digest extract: unsupported type %r", kind)  # path-free
    return ""


__all__ = [
    "MAX_INPUT_BYTES",
    "MAX_PAGES",
    "MAX_CHARS",
    "DigestExtractError",
    "extract_text",
]
